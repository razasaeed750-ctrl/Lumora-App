# app.py — Lumora 🌈
import os, re, io, tempfile
from collections import Counter
from datetime import datetime

import streamlit as st
from pdfminer.high_level import extract_text as extract_text_from_pdf
from docx import Document
from fpdf import FPDF
import plotly.graph_objects as go

# ---------------- CONFIG ----------------
st.set_page_config(page_title="Lumora • Resume Analyzer & Converter", page_icon="🌈", layout="wide")
MAX_FILE_SIZE_MB = 12
ALLOWED = ("pdf", "docx")
CONTACT_WEIGHT = 20
SECTIONS_WEIGHT = 20

# ---------------- THEME / CSS ----------------
st.markdown("""
<style>
:root { --c1:#6a11cb; --c2:#2575fc; --glass: rgba(255,255,255,0.7); }
.stApp {
  background: radial-gradient(1200px 600px at 20% -10%, #ffe6ff 10%, transparent 60%),
              radial-gradient(1200px 700px at 110% 10%, #e9f2ff 10%, transparent 60%),
              linear-gradient(135deg, #fef6ff, #f4f9ff);
}
.hero {
  background: linear-gradient(135deg, var(--c1), var(--c2));
  border-radius: 24px; padding: 42px; color: #fff; text-align:center;
  box-shadow: 0 18px 40px rgba(37,117,252,.25);
}
.hero h1 { font-size: 52px; font-weight: 900; margin: 0; letter-spacing: .4px;}
.hero p { font-size: 18px; opacity: .95; margin: 10px 0 0 0; }
.hero .cta {
  display:inline-block; margin-top: 18px; padding: 14px 26px; border-radius: 999px;
  background:#fff; color:#1e3264; text-decoration:none; font-weight:800;
}
.badge {
  display:inline-flex; align-items:center; gap:.5rem; background:var(--glass); backdrop-filter: blur(8px);
  padding: 8px 14px; border-radius: 999px; font-weight:700; color:#1e3264; box-shadow: 0 2px 10px rgba(0,0,0,.06);
}
.card {
  background: var(--glass); backdrop-filter: blur(8px); border-radius: 18px; padding: 18px 18px;
  box-shadow: 0 8px 24px rgba(0,0,0,.08);
}
.chip {
  display:inline-block; margin: 6px 6px 0 0; padding: 6px 10px; border-radius: 999px; font-size: 13px;
  background: #eef3ff; color:#1e3a8a; border:1px solid #dbe6ff;
}
.kpi {
  background: #fff; border-radius: 16px; padding: 14px; box-shadow: 0 10px 22px rgba(0,0,0,.07); text-align:center;
}
.kpi h3 { margin: 0; font-size: 14px; color:#57606a; font-weight:700; }
.kpi div { font-size: 26px; font-weight: 900; color:#1f2937; }
.section-title { font-weight: 900; font-size: 22px; color:#1e293b; margin-bottom: 6px; }
hr { border: none; height: 1px; background: linear-gradient(90deg, #0000, #e5e7eb, #0000); }
</style>
""", unsafe_allow_html=True)

# ---------------- HELPERS ----------------
def _ext_ok(name:str)->bool:
    return "." in name and name.rsplit(".",1)[1].lower() in ALLOWED

def _read_docx(path_or_file)->str:
    if hasattr(path_or_file, "read"):
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
        tmp.write(path_or_file.read()); tmp.close(); path = tmp.name
    else:
        path = path_or_file
    doc = Document(path)
    text = "\n".join(p.text for p in doc.paragraphs)
    try:
        if 'tmp' in locals(): os.unlink(tmp.name)
    except: pass
    return text

def _read_pdf(path_or_file)->str:
    if hasattr(path_or_file, "read"):
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
        tmp.write(path_or_file.read()); tmp.close(); path = tmp.name
    else:
        path = path_or_file
    text = extract_text_from_pdf(path) or ""
    try:
        if 'tmp' in locals(): os.unlink(tmp.name)
    except: pass
    return text

def extract_plain_text(uploaded_file)->str:
    name = uploaded_file.name.lower()
    if name.endswith(".pdf"):   return _read_pdf(uploaded_file)
    if name.endswith(".docx"):  return _read_docx(uploaded_file)
    return ""

def tokenize(text:str):
    return re.findall(r"\b[a-zA-Z0-9\+#\.\-\_]{2,}\b", text.lower())

def top_keywords_from_jd(jd_text, n=50):
    toks = tokenize(jd_text); from collections import Counter
    c = Counter(toks)
    stop = {"the","and","for","with","that","this","from","your","you","are","a","an","to","in","on","of","or"}
    cand = [(t, c[t]) for t in c if t not in stop and not t.isdigit()]
    cand.sort(key=lambda x: x[1], reverse=True)
    return [t for t,_ in cand][:n]

def score_cv(cv_text:str, jd_text:str|None=None, top_k:int=30):
    s = cv_text.lower()
    tokens_cv = set(tokenize(s))
    jd_tokens = set(top_keywords_from_jd(jd_text, n=top_k)) if jd_text else set()

    if jd_tokens:
        matched = sorted(jd_tokens & tokens_cv)
        kw_score = int(min(100, (len(matched) / max(1,len(jd_tokens))) * 60))
        kw_total = len(jd_tokens)
    else:
        generic = {"python","aws","docker","sql","linux","git","javascript","java","c++","html","css","kubernetes","node","react","data","analysis","machine","learning"}
        matched = sorted(generic & tokens_cv)
        kw_score = int(min(100, (len(matched) / max(1,len(generic))) * 60))
        kw_total = len(generic)

    has_contact = bool(re.search(r"\b[\w\.-]+@[\w\.-]+\.\w{2,}\b", cv_text) or re.search(r"\b\d{7,}\b", cv_text))
    has_sections = any(sec in s for sec in ("experience","education","skills","projects","summary","work"))
    contact_score  = CONTACT_WEIGHT if has_contact else 0
    sections_score = SECTIONS_WEIGHT if has_sections else 0
    wc = len(re.findall(r"\w+", cv_text))
    wc_score = 10 if wc>=300 else (5 if wc>=150 else 0)
    total = min(100, kw_score + contact_score + sections_score + wc_score)

    return {
        "score": total, "kw_score":kw_score, "contact_score":contact_score, "sections_score":sections_score, "wc_score":wc_score,
        "matched": matched, "kw_count": len(matched), "kw_total": kw_total,
        "has_contact":has_contact, "has_sections":has_sections, "word_count":wc
    }

# ---- Document conversion helpers ----
def pdf_to_docx_bytes(pdf_file)->bytes:
    text = _read_pdf(pdf_file)  # plain text extraction
    doc = Document()
    doc.add_heading('Converted from PDF (text-only)', level=1)
    for line in text.splitlines():
        doc.add_paragraph(line)
    buff = io.BytesIO(); doc.save(buff); return buff.getvalue()

def docx_to_pdf_bytes(docx_file)->bytes:
    text = _read_docx(docx_file)
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    # Basic Latin chars best supported; for other scripts embed TTF font with add_font().
    for line in text.splitlines():
        pdf.multi_cell(0, 8, txt=line)
    out = pdf.output(dest="S").encode("latin-1", errors="ignore")
    return out

# ---------------- SIDEBAR NAV ----------------
st.sidebar.markdown("<div class='badge'>🌈 Lumora • Smart Career Tools</div>", unsafe_allow_html=True)
page = st.sidebar.radio("Navigate", ["🏠 Home", "📄 Resume Analyzer", "🔄 Document Converter", "ℹ️ About"])
st.sidebar.caption("v1.0 • Made with ❤️")

# ---------------- HOME ----------------
if page == "🏠 Home":
    st.markdown("""
    <div class='hero'>
      <h1>Lumora</h1>
      <p>Turn your CV into an opportunity magnet — score it like an ATS, get instant fixes, and convert documents in seconds.</p>
      <a class='cta' href='#analyze'>✨ Start Analyzing</a>
    </div>
    """, unsafe_allow_html=True)
    st.write("")
    c1,c2,c3 = st.columns(3)
    with c1: st.markdown("<div class='kpi'><h3>⚡ Instant</h3><div>Real-time scoring</div></div>", unsafe_allow_html=True)
    with c2: st.markdown("<div class='kpi'><h3>🎯 Accurate</h3><div>Keyword matching</div></div>", unsafe_allow_html=True)
    with c3: st.markdown("<div class='kpi'><h3>🔄 Versatile</h3><div>PDF ↔ Word</div></div>", unsafe_allow_html=True)
    st.write("---")
    st.subheader("What you can do")
    st.write("• Analyze and improve your CV with **ATS-style scoring**\n• Extract text from PDFs/DOCX\n• Convert **PDF → Word** and **Word → PDF**\n• Download your results instantly")
    st.balloons()

# ---------------- RESUME ANALYZER ----------------
elif page == "📄 Resume Analyzer":
    st.markdown("<div id='analyze'></div>", unsafe_allow_html=True)
    st.markdown("<h2 class='section-title'>📄 Resume Analyzer</h2>", unsafe_allow_html=True)

    left, right = st.columns([1.7,1.3])
    with left:
        cv_file = st.file_uploader("📂 Upload CV (PDF or DOCX)", type=["pdf","docx"])
        jd = st.text_area("✍️ Paste Job Description / keywords (optional)", height=130,
                          placeholder="Paste the JD here to match keywords more precisely …")
        topk = st.slider("Top K keywords from JD", 5, 60, 30)
        go_btn = st.button("🚀 Analyze CV", type="primary", use_container_width=True)
    with right:
        st.markdown("<div class='card'><b>Tips</b><br>• Use multiple keyword forms (Python, Python3, Django)"
                    "<br>• Add sections: Experience, Education, Skills<br>• Prefer PDF/DOCX with selectable text</div>",
                    unsafe_allow_html=True)

    if go_btn:
        if not cv_file:
            st.error("Please upload a CV file first.")
        else:
            cv_file.seek(0, io.SEEK_END); size_mb = cv_file.tell()/(1024*1024); cv_file.seek(0)
            if size_mb > MAX_FILE_SIZE_MB:
                st.error(f"File too large ({size_mb:.1f} MB). Max {MAX_FILE_SIZE_MB} MB.")
            elif not _ext_ok(cv_file.name):
                st.error("Only PDF or DOCX files are supported.")
            else:
                with st.spinner("Extracting and analyzing…"):
                    text = extract_plain_text(cv_file)
                    if not text.strip():
                        st.warning("No extractable text found (scanned image?). Try an OCR’d PDF.")
                    res = score_cv(text, jd_text=jd if jd.strip() else None, top_k=topk)

                # Gauge (Plotly)
                fig = go.Figure(go.Indicator(
                    mode="gauge+number",
                    value=res["score"],
                    number={'suffix': " / 100", 'font': {'size': 36}},
                    gauge={'axis': {'range': [0, 100]},
                           'bar': {'color': "#2575fc"},
                           'steps': [
                               {'range': [0, 50], 'color': "#ffe5e5"},
                               {'range': [50, 75], 'color': "#fff4cc"},
                               {'range': [75, 100], 'color': "#e5ffe9"}]},
                    title={'text': "ATS Score", 'font': {'size': 20}}
                ))
                st.plotly_chart(fig, use_container_width=True)

                # KPI Row
                k1,k2,k3,k4 = st.columns(4)
                k1.markdown(f"<div class='kpi'><h3>Keywords</h3><div>{res['kw_count']} / {res['kw_total']}</div></div>", unsafe_allow_html=True)
                k2.markdown(f"<div class='kpi'><h3>Contact</h3><div>{'Yes' if res['has_contact'] else 'No'}</div></div>", unsafe_allow_html=True)
                k3.markdown(f"<div class='kpi'><h3>Sections</h3><div>{'OK' if res['has_sections'] else 'Missing'}</div></div>", unsafe_allow_html=True)
                k4.markdown(f"<div class='kpi'><h3>Words</h3><div>{res['word_count']}</div></div>", unsafe_allow_html=True)

                st.write("")
                st.markdown("#### ✅ Matched Keywords")
                if res["matched"]:
                    st.markdown("".join([f"<span class='chip'>{t}</span>" for t in res["matched"]]), unsafe_allow_html=True)
                else:
                    st.caption("No matches found yet — tailor your CV to the JD keywords.")

                st.write(""); st.markdown("#### ✨ Suggestions")
                suggestions = []
                if not res["has_contact"]: suggestions.append("Add **email** and **phone** at the top.")
                if not res["has_sections"]: suggestions.append("Include **Experience**, **Education**, and **Skills** sections.")
                if res["word_count"] < 150: suggestions.append("Add more **bullet points** with **metrics** and outcomes.")
                if res["kw_count"] == 0: suggestions.append("Mirror **keywords** from the JD naturally in your CV.")
                if suggestions:
                    st.markdown("<div class='card'>" + "<br>".join("• " + s for s in suggestions) + "</div>", unsafe_allow_html=True)
                else:
                    st.success("Looks solid! Fine-tune phrasing and quantify achievements.")

                st.download_button("⬇️ Download Extracted Text", text.encode("utf-8"),
                                   file_name="cv_extracted.txt", mime="text/plain", use_container_width=True)
                st.toast("Analysis complete!", icon="✅")

# ---------------- DOCUMENT CONVERTER ----------------
elif page == "🔄 Document Converter":
    st.markdown("<h2 class='section-title'>🔄 Document Converter</h2>", unsafe_allow_html=True)
    tab1, tab2 = st.tabs(["📄 PDF → Word (.docx)", "📝 Word (.docx) → PDF"])

    with tab1:
        up_pdf = st.file_uploader("Upload PDF", type=["pdf"], key="pdf2docx")
        if st.button("Convert to Word (.docx)", use_container_width=True, key="b1"):
            if not up_pdf:
                st.error("Please upload a PDF.")
            else:
                up_pdf.seek(0, io.SEEK_END); size_mb = up_pdf.tell()/(1024*1024); up_pdf.seek(0)
                if size_mb > MAX_FILE_SIZE_MB:
                    st.error(f"File too large ({size_mb:.1f} MB). Max {MAX_FILE_SIZE_MB} MB.")
                else:
                    with st.spinner("Converting…"):
                        docx_bytes = pdf_to_docx_bytes(up_pdf)  # text-only (no original layout)
                    st.success("Done! Download your Word file below.")
                    st.download_button("⬇️ Download .docx",
                        data=docx_bytes, file_name="converted_from_pdf.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True)

    with tab2:
        up_docx = st.file_uploader("Upload Word (.docx)", type=["docx"], key="docx2pdf")
        if st.button("Convert to PDF", use_container_width=True, key="b2"):
            if not up_docx:
                st.error("Please upload a DOCX.")
            else:
                up_docx.seek(0, io.SEEK_END); size_mb = up_docx.tell()/(1024*1024); up_docx.seek(0)
                if size_mb > MAX_FILE_SIZE_MB:
                    st.error(f"File too large ({size_mb:.1f} MB). Max {MAX_FILE_SIZE_MB} MB.")
                else:
                    with st.spinner("Converting…"):
                        pdf_bytes = docx_to_pdf_bytes(up_docx)  # simple text PDF
                    st.success("Done! Download your PDF below.")
                    st.download_button("⬇️ Download .pdf",
                        data=pdf_bytes, file_name="converted_from_docx.pdf",
                        mime="application/pdf", use_container_width=True)

    st.caption("Note: Conversions here are **text-first** (not exact visual layout). For pixel-perfect conversions use a desktop tool or paid API.")

# ---------------- ABOUT ----------------
else:
    st.markdown("<h2 class='section-title'>ℹ️ About Lumora</h2>", unsafe_allow_html=True)
    st.write("""
**Lumora** helps you:
- Analyze CVs with **ATS-style scoring**
- Get **instant suggestions** to improve
- Convert between **PDF** and **Word** quickly

Built with Streamlit • pdfminer.six • python-docx • fpdf2 • Plotly.
    """)
    st.write("— Made for demos, portfolios, and Fiverr gigs ✨")
